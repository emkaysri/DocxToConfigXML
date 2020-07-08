import json, re, ctypes,  os, sys
from docx import Document
import tkinter as tk
from tkinter import filedialog




root = tk.Tk()
root.withdraw()

MessageBox = ctypes.windll.user32.MessageBoxW
MessageBox(None, "Select docx file to open", 'Open file', 0x1000)
file_path = filedialog.askopenfilename()

# determine if application is a script file or frozen exe
if hasattr(sys, 'frozen'):
    basis = sys.executable
else:
    basis = sys.argv[0]

required_folder = os.path.split(basis)[0]

with open(os.path.join(required_folder, 'comments.json')) as json_file:
    comments = json.load(json_file)

directoryPath = os.path.split(file_path)[0] + "/Output-Files/"

if not os.path.exists(directoryPath):
        os.makedirs(directoryPath)    
        
# Load the first table from your document. In your example file,
# there is only one table, so I just grab the first one.
#Document = Document(input(str("File Location:")))

document = Document(file_path)
f = open(directoryPath + "ClientApp.config", "w")
table = document.tables[2]
newVariables = []

# Data will be a list of rows represented as dictionaries
# containing each row's data.
data = []
section = ''
Radio_Onco_Clinical = ''

keys = None
f.write("<Configuration>\n")
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)

    # Establish the mapping based on the first row
    # headers; these will become the keys of our dictionary
    if i == 0:
        keys = tuple(text)
        continue

    # Construct a dictionary for this row, mapping
    # keys to values for this row
    row_data = dict(zip(keys, text))
    values = list(row_data.values())
    if(all(x == values[0] for x in values)):
        header_text = re.sub(r'\W+', '', values[0])
        if("General" in header_text):
            if(Radio_Onco_Clinical != ''):
                f.write("\t\t</" + section + ">\n")
            section = "General"
            Radio_Onco_Clinical = header_text[header_text.find("-") + 1:]
            f.write("\t\t<" + section + ">\n")
        else:
            f.write("\t\t</" + section + ">\n")
            section = header_text
            f.write("\t\t<" + section + ">\n")
    else:
        value = values[2]
        name = values[0]
        for i in ['<','>',' ','\n']:
            name = name.replace(i, '') 
        if ("empty" in value.lower()):
            value = ""
        if ("note" in value.lower()):
            value = ""            
        for i in ['<','>',' ','\n']:
            value = value.replace(i, '') 
        
        try:
            f.write("\t\t\t<" + name + ">" + value + "</" + name + ">\t" + comments[name] + "\n")
        except KeyError:
            f.write("\t\t\t<" + name + ">" + value + "</" + name + ">\n")
            newVariables.append(name)
            
if (len(newVariables) > 0):
    MessageBox(None, "Please add the following variables: \n" + '\n'.join(map(str, newVariables)) +"\n\nCheck addvariables.txt for the list of variables", 'Add variables to ClientConfig.cs', 0x1000)
    z = open(directoryPath + "addvariables.txt", "w")
    z.write('\n'.join(map(str, newVariables)))
    z.close()
    
f.write("\t\t</" + section + ">\n")
f.write("\t</" + Radio_Onco_Clinical + ">\n")
f.write("</Configuration>")
f.close()

MessageBox(None, "Completed Successfully", 'Success', 0x1000)